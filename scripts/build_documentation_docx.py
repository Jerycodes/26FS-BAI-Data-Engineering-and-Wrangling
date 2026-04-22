"""
build_documentation_docx.py - Erzeugt DOKUMENTATION.docx aus DOKUMENTATION.md.

Konvertiert das Markdown-Dokument in ein Word-File für die Abgabe. Überschriften,
Listen, Tabellen und Code-Blöcke werden in native Word-Elemente umgewandelt.

Verwendung:
    python scripts/build_documentation_docx.py
"""

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


SRC = Path("DOKUMENTATION.md")
DST = Path("DOKUMENTATION.docx")


def parse_inline_formatting(run_paragraph, text: str) -> None:
    """Parst **bold**, `code` und normaler Text; fuegt Runs dem Paragraph hinzu."""
    pattern = re.compile(r"(\*\*.+?\*\*|`[^`]+`)")
    parts = pattern.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = run_paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = run_paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        else:
            run_paragraph.add_run(part)


def add_heading(doc: Document, text: str, level: int) -> None:
    heading = doc.add_heading(level=min(level, 4))
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    parse_inline_formatting(heading, text)


def add_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    parse_inline_formatting(p, text)


def add_code_block(doc: Document, lines: list[str]) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run("\n".join(lines))
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Light Grid Accent 1"
    for row_idx, row in enumerate(rows):
        for col_idx, cell_text in enumerate(row):
            cell = table.cell(row_idx, col_idx)
            cell.text = ""
            p = cell.paragraphs[0]
            parse_inline_formatting(p, cell_text.strip())
            if row_idx == 0:
                for run in p.runs:
                    run.bold = True
    doc.add_paragraph()


def main() -> None:
    if not SRC.exists():
        raise SystemExit(f"Quelle fehlt: {SRC}")

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    lines = SRC.read_text(encoding="utf-8").splitlines()

    i = 0
    list_buffer: list[tuple[str, str]] = []

    def flush_list() -> None:
        nonlocal list_buffer
        for kind, text in list_buffer:
            p = doc.add_paragraph(style="List Bullet" if kind == "-" else "List Number")
            parse_inline_formatting(p, text)
        list_buffer = []

    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()

        if not line.strip():
            flush_list()
            i += 1
            continue

        if line.startswith("```"):
            flush_list()
            code_lines: list[str] = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                code_lines.append(lines[i])
                i += 1
            add_code_block(doc, code_lines)
            i += 1
            continue

        if line.startswith("#"):
            flush_list()
            m = re.match(r"^(#+)\s+(.*)$", line)
            if m:
                add_heading(doc, m.group(2), len(m.group(1)))
            i += 1
            continue

        if line.strip() == "---":
            flush_list()
            i += 1
            continue

        if line.startswith("|") and line.endswith("|"):
            flush_list()
            table_rows: list[list[str]] = []
            while i < len(lines) and lines[i].startswith("|"):
                content = lines[i].strip().strip("|")
                if re.match(r"^\s*[-:]+\s*\|", lines[i].strip().strip("|") + "|"):
                    # Trennzeile
                    i += 1
                    continue
                cells = [c.strip() for c in content.split("|")]
                # Heuristik: Trennzeilen bestehen nur aus ---
                if all(re.match(r"^[-:\s]+$", c) for c in cells) and cells:
                    i += 1
                    continue
                table_rows.append(cells)
                i += 1
            add_table(doc, table_rows)
            continue

        m_list = re.match(r"^(\s*)[-*]\s+(.*)$", line)
        if m_list:
            list_buffer.append(("-", m_list.group(2)))
            i += 1
            continue

        m_num = re.match(r"^(\s*)\d+\.\s+(.*)$", line)
        if m_num:
            list_buffer.append(("1", m_num.group(2)))
            i += 1
            continue

        flush_list()
        m_quote = re.match(r"^>\s*(.*)$", line)
        if m_quote:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            run = p.add_run(m_quote.group(1))
            run.italic = True
        else:
            add_paragraph(doc, line)
        i += 1

    flush_list()

    doc.save(DST)
    size_kb = DST.stat().st_size / 1024
    print(f"Gespeichert: {DST} ({size_kb:.1f} KB)")


if __name__ == "__main__":
    main()
