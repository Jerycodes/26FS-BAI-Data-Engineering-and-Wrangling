"""
build_report_docx.py - Erzeugt den abgabefertigen Bericht DOKUMENTATION.docx.

Baut das Word-Dokument direkt mit python-docx (nicht über Markdown), damit
Seitenumbrüche, Inhaltsverzeichnis, Abbildungs- und Tabellenverzeichnis, farbliche
Tabellen-Markierungen, Glossar und Theorie-Anhang sauber möglich sind.

Die eingebetteten Abbildungen erzeugt zuvor scripts/build_report_figures.py.

Aufruf vom Projekt-Root:
    python scripts/build_report_docx.py
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Cm, RGBColor

DST = Path("DOKUMENTATION.docx")
FIGDIR = Path("docs/figures")
PIPELINE_IMG = Path("docs/architektur/pipeline.png")

COL_REAL = "DCE6F1"   # helles Blau  = im Hauptpfad der Analyse verwendet
COL_POC = "FCE4D6"    # helles Orange = eigene Umsetzung von Nachrichten/Sentiment
COL_HEAD = "404040"   # Kopfzeile dunkelgrau

# Verzeichnisse (vorne im Dokument, Nummern passen zur Reihenfolge im Text)
FIGURE_LIST = [
    "Abbildung 1: Aufbau der Datenpipeline von der Quelle bis zu den Auswertungen",
    "Abbildung 2: Verteilung der Tagesveränderungen mit Boxplot der Extremwerte",
    "Abbildung 3: Quellenvergleich vor und nach der Datums-Ausrichtung",
    "Abbildung 4: Eigenes Sentiment im Vergleich zum vorberechneten Sentiment von EODHD",
    "Abbildung 5: Zusammenhang zwischen Sentiment und Kursveränderung je zeitlicher Verschiebung",
    "Abbildung 6: EUR/USD, Kursniveau und wöchentliches Nachrichten-Sentiment im Zeitverlauf",
    "Abbildung 7: GBP/USD, Kursniveau und wöchentliches Nachrichten-Sentiment im Zeitverlauf",
    "Abbildung 8: Sentiment einer Woche und kumulative Kursbewegung der Folgewochen",
    "Abbildung 9: Ölpreise WTI und Brent im Zeitverlauf (Anhang E)",
    "Abbildung 10: Monatliche Datenabdeckung je Quelle und Währungspaar (Anhang E)",
]
TABLE_LIST = [
    "Tabelle 1: Eingesetzte Werkzeuge nach Aufgabe",
    "Tabelle 2: Übersicht der Datenquellen und geladenen Symbole",
    "Tabelle 3: Umfang, fehlende Werte und Duplikate je Datensatz",
    "Tabelle 4: Übereinstimmung von Yahoo und EODHD vor der Korrektur",
    "Tabelle 5: Zusammenhang am selben Tag mit und ohne Interpolation",
    "Tabelle 6: Stärkster Zusammenhang je Wechselkurs auf Tagesebene",
    "Tabelle 7: Sentiment einer Woche und Kursbewegung in den Folgewochen",
    "Tabelle 8: Zusammenhang im selben Zeitraum je Zeithorizont",
    "Tabelle 9: Granger-Test, p-Werte für den Vorhersagebeitrag des Sentiments",
    "Tabelle 10: Glossar der wichtigsten Begriffe",
    "Tabelle 11: Artikel pro Tag und Währungspaar bei EODHD (Anhang E)",
]


# ---------------------------------------------------------------------------
# Hilfsfunktionen
# ---------------------------------------------------------------------------

def shade_cell(cell, fill_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def set_cell_text(cell, text, *, bold=False, white=False, size=10):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if white:
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(9)


def add_table(doc, table_caption, header, rows, *, row_colors=None, widths=None, font=9):
    caption(doc, table_caption)
    t = doc.add_table(rows=1, cols=len(header))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    for j, h in enumerate(header):
        c = t.rows[0].cells[j]
        set_cell_text(c, h, bold=True, white=True, size=font)
        shade_cell(c, COL_HEAD)
    for i, row in enumerate(rows):
        cells = t.add_row().cells
        for j, val in enumerate(row):
            set_cell_text(cells[j], str(val), size=font)
            if row_colors and row_colors[i]:
                shade_cell(cells[j], row_colors[i])
    if widths:
        for row in t.rows:
            for j, w in enumerate(widths):
                row.cells[j].width = Cm(w)
    doc.add_paragraph()
    return t


def add_figure(doc, img_path, fig_caption, width=15.0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if Path(img_path).exists():
        p.add_run().add_picture(str(img_path), width=Cm(width))
    else:
        r = p.add_run(f"[Abbildung fehlt: {img_path}]")
        r.italic = True
    caption(doc, fig_caption)
    doc.add_paragraph()


def add_legend(doc):
    t = doc.add_table(rows=1, cols=2)
    t.style = "Table Grid"
    set_cell_text(t.rows[0].cells[0], "Im Hauptpfad der Analyse verwendet", size=9)
    shade_cell(t.rows[0].cells[0], COL_REAL)
    set_cell_text(t.rows[0].cells[1], "Eigene Umsetzung von Nachrichten und Sentiment (Demonstration)", size=9)
    shade_cell(t.rows[0].cells[1], COL_POC)
    doc.add_paragraph()


def h1(doc, text):
    doc.add_page_break()
    doc.add_heading(text, level=1)


def h2(doc, text):
    doc.add_heading(text, level=2)


def h3(doc, text):
    doc.add_heading(text, level=3)


def para(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run(text)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run(text)
    return p


def label_block(doc, label, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(label)
    r.bold = True
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p2.paragraph_format.left_indent = Cm(0.6)
    p2.add_run(text)


def add_toc(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    ph = OxmlElement("w:t")
    ph.text = "Rechtsklick auf dieses Feld, dann Feld aktualisieren, um das Inhaltsverzeichnis zu füllen."
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for el in (begin, instr, sep, ph, end):
        run._r.append(el)


def update_fields_on_open(doc):
    el = OxmlElement("w:updateFields")
    el.set(qn("w:val"), "true")
    doc.settings.element.append(el)


# ---------------------------------------------------------------------------
# Dokument
# ---------------------------------------------------------------------------

def build():
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    # Schmalere Ränder: mehr Platz für Text und Abbildungen (A4: 21 cm breit)
    for section in doc.sections:
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)

    # ---- Titelseite ----
    for _ in range(4):
        doc.add_paragraph()
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Zusammenhang zwischen Nachrichten-Sentiment und Wechselkursen"); r.bold = True; r.font.size = Pt(22)
    st = doc.add_paragraph(); st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = st.add_run("Ein Data-Engineering- und Wrangling-Projekt zu EUR/USD, EUR/CHF und GBP/USD"); r.font.size = Pt(13)
    for _ in range(6):
        doc.add_paragraph()
    for line in [
        "Modul: Data Engineering and Wrangling",
        "Studiengang Business Artificial Intelligence (BAI)",
        "Fachhochschule Nordwestschweiz (FHNW)",
        "Frühjahrssemester 2026",
        "",
        "Verfasst von:",
        "Jeremy Nathan, Stirling Mulholland, Fenlin Chirakkal",
    ]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(line).font.size = Pt(12)

    # ---- Inhaltsverzeichnis ----
    doc.add_page_break()
    doc.add_heading("Inhaltsverzeichnis", level=1)
    add_toc(doc)

    # ---- Abbildungs- und Tabellenverzeichnis ----
    doc.add_page_break()
    doc.add_heading("Abbildungsverzeichnis", level=1)
    for line in FIGURE_LIST:
        para(doc, line)
    doc.add_heading("Tabellenverzeichnis", level=1)
    for line in TABLE_LIST:
        para(doc, line)

    # =====================================================================
    h1(doc, "1. Einleitung")

    h2(doc, "1.1 Ausgangslage und Ziel")
    para(doc, "Wechselkurse reagieren laufend auf neue Informationen, und ein grosser Teil dieser "
              "Informationen erreicht den Markt in Form von Finanznachrichten. Dieses Projekt untersucht, "
              "ob sich zwischen der Stimmung in solchen Nachrichten und der Bewegung von Wechselkursen ein "
              "messbarer Zusammenhang zeigt. Betrachtet werden die drei Wechselkurse Euro zu US-Dollar "
              "(EUR/USD), Euro zu Schweizer Franken (EUR/CHF) und Britisches Pfund zu US-Dollar (GBP/USD) "
              "im Zeitraum vom 1. Januar 2022 bis zum 22. April 2026.")
    para(doc, "Der Schwerpunkt der Arbeit liegt auf der Aufbereitung der Daten: aus mehreren, "
              "unterschiedlich aufgebauten Quellen vergleichbare Zeitreihen bauen, Qualitätsprobleme "
              "erkennen und beheben, und jede Entscheidung nachvollziehbar begründen. Die Beantwortung der "
              "Forschungsfrage ist das Anwendungsbeispiel, an dem wir die aufbereiteten Daten zeigen.")

    h2(doc, "1.2 Forschungsfrage")
    para(doc, "Die Arbeit geht der folgenden Frage nach:")
    q = doc.add_paragraph(); q.paragraph_format.left_indent = Cm(0.6)
    q.add_run("Hat die Stimmung von Finanznachrichten einen Einfluss auf Wechselkurse, und falls ja, mit "
              "welcher zeitlichen Verzögerung?").italic = True
    para(doc, "Anschaulich formuliert: Wenn die Nachrichten in einem bestimmten Zeitraum überwiegend "
              "negativ waren, bewegt sich der Wechselkurs dann im selben Zeitraum oder erst danach nach "
              "unten? Der betrachtete Zeitraum ist dabei frei wählbar. Wir verwenden im Bericht meist eine "
              "Woche als anschauliches Beispiel, prüfen den Zusammenhang aber ebenso auf Tagesebene und auf "
              "anderen Zeithorizonten wie zehn Tagen oder einem Monat.")
    para(doc, "Unter Stimmung verstehen wir das Sentiment: eine Zahl, die ausdrückt, wie negativ oder "
              "positiv ein Text formuliert ist. Sie reicht von minus eins (sehr negativ) über null "
              "(neutral) bis plus eins (sehr positiv).")

    h2(doc, "1.3 Hypothesen")
    para(doc, "Wir übersetzen die Frage in drei mögliche Antworten. Dabei steht die zeitliche "
              "Verschiebung k für den Abstand zwischen Nachrichten und Kursbewegung, gemessen in der "
              "jeweils betrachteten Zeiteinheit (Tage, Wochen oder Monate).")
    label_block(doc, "Hypothese H1: das Sentiment läuft dem Kurs voraus",
                "Die Stimmung eines Zeitraums sagt die Kursbewegung der folgenden Zeiträume voraus. "
                "Negative Nachrichten heute bedeuten einen fallenden Kurs in den nächsten Tagen oder "
                "Wochen. Diese Hypothese würde einen Vorlauf des Sentiments behaupten.")
    label_block(doc, "Alternative A1: Sentiment und Kurs bewegen sich gleichzeitig",
                "Stimmung und Kurs ändern sich im selben Zeitraum, ohne dass eine Reihe der anderen "
                "vorausläuft. Das Sentiment ist dann eine Begleitinformation, kein Frühindikator.")
    label_block(doc, "Alternative A2: der Kurs läuft dem Sentiment voraus",
                "Der Kurs bewegt sich zuerst, die Nachrichten berichten erst danach darüber. Die "
                "zeitliche Reihenfolge wäre dann umgekehrt.")

    h2(doc, "1.4 Aufbau der Arbeit")
    para(doc, "Die Daten durchlaufen mehrere Schritte, die zugleich den Aufbau dieses Berichts bilden: "
              "Daten laden (Kapitel 4), bereinigen und auf Qualität prüfen (Kapitel 5), harmonisieren "
              "(Kapitel 6), in Analysegrössen umrechnen (Kapitel 7), das Sentiment bestimmen (Kapitel 8) "
              "und schliesslich auswerten (Kapitel 9). Abbildung 1 zeigt diesen Weg im Überblick. Die "
              "Rohdaten bleiben dabei unverändert erhalten, abgeleitete Ergebnisse entstehen in einer "
              "getrennten Verarbeitungsschicht. So lässt sich jeder Schritt wiederholen und prüfen, ohne "
              "die Originaldaten zu verändern, und ein Fehler in der Verarbeitung erfordert kein neues "
              "Laden von den Anbietern. Jeder Verarbeitungsschritt ist als eigenes Skript umgesetzt und "
              "liefert bei gleichen Eingaben dasselbe Ergebnis (Anhang C zeigt die Befehle).")
    add_figure(doc, PIPELINE_IMG, "Abbildung 1: Aufbau der Datenpipeline von der Quelle bis zu den Auswertungen.", width=17)
    para(doc, "Das Projekt besteht aus zwei Teilen. Den Kern bildet die Datenpipeline mit den "
              "Wechselkursen von Yahoo Finance, EODHD und MetaTrader 5 sowie den Finanznachrichten von "
              "EODHD samt deren fertig berechnetem Sentiment. Darauf beruht die Analyse.")
    para(doc, "Speziell für die Nachrichten und deren Stimmungsanalyse zeigen wir zusätzlich, wie sich "
              "dieser Schritt selbst umsetzen lässt, und zwar auf zwei Arten. Erstens nehmen wir den "
              "Nachrichtentext von EODHD und berechnen das Sentiment mit einem eigenen Verfahren, statt "
              "den fertigen Wert von EODHD zu übernehmen. Zweitens beschaffen wir die Nachrichten "
              "vollständig selbst über Webscraping und berechnen auch dort das Sentiment selbst. So decken "
              "wir den gesamten Weg von der eigenen Datenbeschaffung bis zum fertigen Stimmungswert ab.")

    # =====================================================================
    h1(doc, "2. Eingesetzte Werkzeuge")
    para(doc, "Das Projekt ist in Python umgesetzt. Die folgende Tabelle zeigt, welche Programmbibliothek "
              "wir für welche Aufgabe verwendet haben.")
    add_table(
        doc, "Tabelle 1: Eingesetzte Werkzeuge nach Aufgabe",
        ["Aufgabe", "Werkzeug"],
        [
            ["Programmiersprache", "Python 3.12"],
            ["Wechselkurse und Ölpreise laden", "yfinance"],
            ["EODHD-Daten laden (Kurse und Nachrichten)", "requests"],
            ["Nachrichten sammeln (Webscraping)", "requests, feedparser, BeautifulSoup"],
            ["Datenverarbeitung und Berechnung", "pandas, numpy, scipy"],
            ["Eigene Sentiment-Analyse", "TextBlob"],
            ["Statistische Tests (Granger)", "statsmodels"],
            ["Abbildungen im Bericht", "matplotlib, seaborn"],
            ["Interaktive Darstellung (Anhang D)", "plotly, streamlit"],
            ["Pipeline-Diagramm", "Graphviz"],
        ],
        widths=[7.5, 8.5],
    )

    # =====================================================================
    h1(doc, "3. Datengrundlage")

    h2(doc, "3.1 Übersicht der Quellen")
    para(doc, "Wir verwenden bewusst mehrere Quellen pro Datentyp. Nur so lassen sich die Quellen "
              "gegeneinander prüfen, was im Projekt einen echten Datenfehler aufgedeckt hat (Kapitel 5.5). "
              "Tabelle 2 zeigt, welche Quelle welche Wechselkurse und Nachrichten liefert. Blau markiert "
              "ist, was direkt in die Analyse einfliesst, orange die eigene Umsetzung der Nachrichten und "
              "des Sentiments.")
    add_table(
        doc, "Tabelle 2: Übersicht der Datenquellen und geladenen Symbole",
        ["Datentyp", "Quelle", "Geladene Symbole oder Kurse", "Zugang"],
        [
            ["Wechselkurs", "Yahoo Finance", "EURUSD=X, EURCHF=X, GBPUSD=X", "Bibliothek yfinance, kein Login"],
            ["Wechselkurs", "EODHD", "EURUSD.FOREX, EURCHF.FOREX, GBPUSD.FOREX", "REST-Schnittstelle mit Schlüssel"],
            ["Wechselkurs", "MetaTrader 5", "EURUSD (Tages- und 15-Minuten-Daten)", "CSV-Export aus dem Broker ActivTrades"],
            ["Nachrichten", "EODHD", "EURUSD.FOREX, EURCHF.FOREX, GBPUSD.FOREX", "REST-Schnittstelle mit Schlüssel"],
            ["Ölpreis", "Yahoo Finance", "CL=F (WTI), BZ=F (Brent)", "Bibliothek yfinance, kein Login"],
            ["Nachrichten", "RSS-Feeds (4 aktive)", "ForexLive, FXStreet, Yahoo Finance, Google News", "requests und feedparser"],
            ["Nachrichten", "Reddit", "r/Forex, r/investing, r/economics", "öffentlicher JSON-Endpunkt"],
        ],
        row_colors=[COL_REAL, COL_REAL, COL_REAL, COL_REAL, COL_REAL, COL_POC, COL_POC],
        widths=[2.4, 3.0, 5.8, 4.4],
    )
    add_legend(doc)

    h2(doc, "3.2 Warum diese Quellen")
    label_block(doc, "EODHD",
                "Der Ausgangspunkt für EODHD war, dass dieser Anbieter Finanznachrichten direkt zusammen "
                "mit einem fertig berechneten Stimmungswert liefert. Weil EODHD die Nachrichten "
                "bereitstellt, lag es nahe, auch die Wechselkurse von derselben Quelle zu beziehen. So "
                "stammen Kurs und Nachricht aus demselben System.")
    label_block(doc, "Yahoo Finance",
                "Yahoo Finance dient als zweite, kostenlose und unabhängige Quelle für dieselben "
                "Wechselkurse. Mit einer zweiten Quelle lässt sich prüfen, ob die Werte von EODHD "
                "plausibel sind.")
    label_block(doc, "MetaTrader 5",
                "MetaTrader 5 stammt direkt von einem handelbaren Broker. Bei Yahoo Finance und EODHD "
                "handelt man nicht, es sind reine Datenanbieter. Wenn sich Quellen unterscheiden, ist es "
                "sinnvoll, eine Quelle einzubeziehen, die von dem Ort kommt, an dem auch tatsächlich "
                "gehandelt wird. Die Daten stammen aus einem Demokonto beim Broker ActivTrades, über den "
                "sich MetaTrader 5 laden lässt. MetaTrader 5 dient zum einen dem Quellenvergleich und "
                "fliesst zum anderen bei EUR/USD als dritte Quelle in den Tages-Mittelwert ein "
                "(Kapitel 7.3), bis zu seinem Exportende am 26. Dezember 2025.")
    label_block(doc, "Ölpreise",
                "Der Ölpreis gilt als häufig genannter Einflussfaktor auf rohstoffnahe Währungen und "
                "dient als zuschaltbare Vergleichsreihe.")

    # =====================================================================
    h1(doc, "4. Daten laden")
    para(doc, "Jede Quelle hat ein eigenes Lade-Skript im Ordner src/data_loading. Alle folgen demselben "
              "Muster: Daten holen und unverändert im Ordner data/raw ablegen. In diesem Schritt findet "
              "bewusst keine Bereinigung statt, damit das Rohmaterial erhalten bleibt.")

    h2(doc, "4.1 Wechselkurse")
    para(doc, "Von Yahoo Finance laden wir die drei Symbole EURUSD=X, EURCHF=X und GBPUSD=X als Tagesdaten "
              "über die Bibliothek yfinance. Ein Login ist nicht nötig.")
    para(doc, "Von EODHD laden wir die drei Symbole EURUSD.FOREX, EURCHF.FOREX und GBPUSD.FOREX über den "
              "Aufruf https://eodhd.com/api/eod/{Symbol} mit dem Parameter period=d für Tagesdaten. Alle "
              "Aufrufe an EODHD, sowohl für Kurse als auch für Nachrichten, nutzen denselben einen "
              "Zugangsschlüssel aus der lokalen Datei .env.")
    para(doc, "Von MetaTrader 5 exportieren wir EURUSD als Tagesdaten und zusätzlich als 15-Minuten-Daten. "
              "Der Export erfolgt manuell aus der Handelsplattform.")

    h2(doc, "4.2 Nachrichten von EODHD")
    para(doc, "Die Nachrichten laden wir über den Aufruf https://eodhd.com/api/news. Gefiltert wird je "
              "Währungspaar über den Symbol-Parameter s, mit den drei Werten EURUSD.FOREX, EURCHF.FOREX "
              "und GBPUSD.FOREX. Pro Aufruf holen wir bis zu 1000 Artikel (Parameter limit=1000) und "
              "blättern über den Parameter offset weiter, bis keine weiteren Artikel mehr kommen. Pro Paar "
              "speichern wir die Rohdaten als JSON-Datei und zusätzlich eine aufbereitete CSV-Tabelle. Den "
              "verschachtelten Stimmungswert flachen wir dabei in eigene Spalten auf.")

    h2(doc, "4.3 Ölpreise")
    para(doc, "Von Yahoo Finance laden wir die zwei Symbole CL=F (Rohölsorte West Texas Intermediate, kurz "
              "WTI) und BZ=F (Rohölsorte Brent) als Tagesdaten.")

    h2(doc, "4.4 Eigene Nachrichten über Webscraping")
    para(doc, "Für die eigene Datenbeschaffung sammeln wir Nachrichten aus zwei Arten von Quellen. Erstens "
              "aus RSS-Feeds, einem standardisierten Format, über das Webseiten ihre neuesten Artikel "
              "bereitstellen. Wir fragen die folgenden Feeds ab:")
    bullet(doc, "ForexLive: https://www.forexlive.com/feed")
    bullet(doc, "FXStreet: https://www.fxstreet.com/rss/news")
    bullet(doc, "Yahoo Finance: https://finance.yahoo.com/news/rssindex")
    bullet(doc, "Google News (Suche nach forex EUR USD): https://news.google.com/rss/search?q=forex+EUR+USD")
    para(doc, "Ein fünfter Feed, DailyFX, wurde ebenfalls abgefragt, lieferte aber den Fehlercode HTTP 403 "
              "(Zugriff verweigert) und damit keine Daten. Es blieben vier nutzbare Feeds.")
    para(doc, "Zweitens lesen wir die öffentlich verfügbaren Beiträge der drei Reddit-Foren r/Forex "
              "(Sortierungen hot und new), r/investing (Sortierung hot) und r/economics (Sortierung hot) "
              "über deren JSON-Endpunkt. Beim Abfragen der RSS-Feeds holen wir den Inhalt zuerst mit der "
              "Bibliothek requests und übergeben ihn dann an die Bibliothek feedparser; ein direkter "
              "Zugriff scheiterte an der Zertifikatsprüfung.")

    # =====================================================================
    h1(doc, "5. Datenbereinigung und Qualitätsprüfung")

    h2(doc, "5.1 Explorative Übersicht der Kursdaten")
    para(doc, "Vor der Bereinigung haben wir uns einen Überblick über die geladenen Daten verschafft. "
              "Je Währungspaar liegen 1117 Handelstage vor, von Anfang Januar 2022 bis zum 21. April "
              "2026. Für die Prüfung der Verteilung betrachten wir die Tagesveränderung in Prozent, also "
              "die Veränderung des Kurses von einem Handelstag zum nächsten. Abbildung 2 zeigt links die "
              "Verteilung der Tagesveränderungen und rechts den Boxplot mit den Extremwerten.")
    add_figure(doc, FIGDIR / "fig_eda_returns.png",
               "Abbildung 2: Verteilung der Tagesveränderungen mit Boxplot der Extremwerte.", width=16)
    para(doc, "Die Verteilungen sind um null zentriert und symmetrisch, mit einer typischen "
              "Tagesschwankung von 0.46 Prozent bei EUR/USD, 0.55 Prozent bei GBP/USD und 0.35 Prozent "
              "bei EUR/CHF. EUR/CHF schwankt also am wenigsten, was zur Rolle des Schweizer Frankens als "
              "stabiler Währung passt. Auffällig sind die im Vergleich zur Glockenkurve häufigeren "
              "Extremwerte an den Rändern; mit diesen befasst sich Kapitel 5.4.")

    h2(doc, "5.2 Fehlende Werte")
    para(doc, "Beim Umgang mit fehlenden Werten haben wir für jeden Datentyp einzeln entschieden. Der "
              "Grundsatz: Wir füllen eine Lücke nur dann auf, wenn sie rein technisch entsteht, nicht aber "
              "dann, wenn das Fehlen selbst eine Information ist.")
    h3(doc, "Wechselkurse")
    para(doc, "Der Devisenmarkt schliesst am Freitagabend und öffnet erst am Sonntagabend wieder. Yahoo "
              "und MetaTrader liefern darum nur Werktage. EODHD liefert zusätzlich praktisch jeden "
              "Sonntag einen Wert, weil der Markt am Sonntagabend in Asien öffnet; in den unkorrigierten "
              "EODHD-Rohdaten sind ausserdem einzelne Zeilen auf Samstag datiert, ein weiteres Symptom "
              "des in Kapitel 5.5 beschriebenen Datierungsproblems. Diese Wochenendwerte bleiben in den "
              "Rohdaten erhalten. Für die kombinierte Analysereihe verwenden wir aber den "
              "Handelskalender von Yahoo, also Montag bis Freitag. Der Grund: Am Sonntag hätte nur eine "
              "einzige Quelle einen Wert, und der Mittelwert über die Quellen würde an diesen Tagen die "
              "Grundlage wechseln. An Feiertagen fehlen einzelne Werte, die wir nicht erfinden. Ob ein "
              "Auffüllen fehlender Tage durch Interpolation sinnvoll ist, haben wir zusätzlich getestet; "
              "den Vergleich mit und ohne Interpolation zeigen wir in Kapitel 7.4.")
    h3(doc, "Nachrichten-Sentiment")
    para(doc, "Liegt an einem Tag kein Artikel vor, gibt es auch keinen Stimmungswert, und das Feld bleibt "
              "leer. Solche Lücken füllen wir bewusst nicht auf. Wir haben sie als fehlend aufgrund des "
              "Ereignisses selbst eingestuft: Ob ein Artikel existiert, hängt direkt davon ab, ob an "
              "diesem Tag etwas Berichtenswertes passiert ist. Ein Tag ohne Artikel bedeutet also nicht "
              "dasselbe wie ein neutraler Tag mit Stimmung null. Würden wir die Lücke mit null oder mit "
              "dem Nachbarwert füllen, täuschten wir eine Nachrichtenlage vor, die es nicht gab. Die genaue "
              "Einordnung der Arten fehlender Werte steht im Anhang B.")
    para(doc, "Bei der späteren Zusammenfassung auf Wochen oder Monate werden fehlende Tage automatisch "
              "übersprungen. Eine Woche mit nur zwei statt fünf Nachrichtentagen wird also weiterhin "
              "ausgewertet, eben auf Basis der vorhandenen Tage.")
    h3(doc, "Sonderfall EUR/CHF-Nachrichten")
    para(doc, "Für EUR/CHF lieferte EODHD nur sehr wenige Artikel (rund 13 im gesamten Zeitraum). Wir "
              "vermuten, dass dieses Paar international weniger Medienaufmerksamkeit bekommt. So kam es von "
              "EODHD, und wegen der zu dünnen Grundlage haben wir EUR/CHF in der Sentiment-Analyse nicht "
              "verwendet.")

    h2(doc, "5.3 Duplikate")
    para(doc, "Bei den Wechselkursen haben wir geprüft, ob ein Datum je Quelle doppelt vorkommt. Das war "
              "praktisch nicht der Fall: In keiner der Yahoo- und EODHD-Rohdateien gab es doppelte "
              "Datumseinträge. Eine Entfernung von Doppeln war also nicht nötig, der Prüfschritt blieb ohne "
              "Treffer.")
    para(doc, "Bei den selbst gesammelten Nachrichten traten Duplikate dagegen tatsächlich auf. Wir haben "
              "die Feeds zu sechs verschiedenen Zeitpunkten abgefragt. Derselbe Artikel erscheint dabei in "
              "mehreren Abfragen, solange er im Feed steht. Von insgesamt 3012 gesammelten Artikeln waren "
              "541 solche Mehrfachvorkommen. Wir haben sie über die eindeutige Internetadresse des Artikels "
              "(den Link) entfernt und je Artikel das erste Vorkommen behalten.")
    para(doc, "Tabelle 3 fasst Umfang, fehlende Werte und Duplikate aller Datensätze zusammen, mit dem "
              "jeweils festgelegten Umgang.")
    add_table(
        doc, "Tabelle 3: Umfang, fehlende Werte und Duplikate je Datensatz",
        ["Datensatz", "Umfang", "Fehlende Werte", "Duplikate", "Umgang"],
        [
            ["Kurse Yahoo (je Paar)", "1117 Tage, 03.01.2022 bis 21.04.2026",
             "5 Werktage (0.4 Prozent)", "0", "Lücken bleiben leer"],
            ["Kurse EODHD (drei Paare)", "1412 / 1415 / 1175 Tage",
             "keine fehlenden Werktage, zusätzlich Wochenendwerte", "0",
             "Wochenende nur in den Rohdaten (Kapitel 5.2)"],
            ["Kurse MetaTrader (EUR/USD)", "1037 Tage, bis 26.12.2025",
             "3 Werktage (0.3 Prozent)", "0", "dritte Quelle im EUR/USD-Mittel"],
            ["Nachrichten EODHD (EUR/USD, GBP/USD, EUR/CHF)", "28045 / 18172 / 13 Artikel",
             "108 / 128 / 1 Artikel ohne Sentiment", "0 doppelte Links", "Artikel ohne Sentiment bleiben leer"],
            ["Nachrichten Webscraping", "3012 Artikel aus 6 Abfragen",
             "431 Artikel ohne brauchbares Datum", "541 Mehrfachvorkommen", "Duplikate entfernt, 2471 verbleiben"],
        ],
        row_colors=[COL_REAL, COL_REAL, COL_REAL, COL_REAL, COL_POC],
        widths=[3.6, 3.4, 3.6, 2.6, 3.6],
        font=8,
    )
    para(doc, "Die 431 Artikel ohne brauchbares Datum stammen aus Feeds, die kein maschinenlesbares "
              "Veröffentlichungsdatum mitliefern. Sie bleiben im Artikelbestand erhalten, fliessen aber "
              "nicht in die Auswertung auf Tagesebene ein, weil sie sich keinem Tag zuordnen lassen.")

    h2(doc, "5.4 Ausreisser")
    para(doc, "Ausreisser sind einzelne Werte, die weit ausserhalb des üblichen Bereichs liegen. Wir "
              "haben die Tagesveränderungen der drei Wechselkurse mit zwei gängigen Verfahren geprüft, "
              "der z-Wert-Regel und der Interquartilsabstand-Regel (beide im Anhang B erklärt). Die "
              "strengere z-Wert-Regel findet je Wechselkurs 13 bis 14 auffällige Tage, die breiter "
              "gefasste Interquartilsabstand-Regel 39 bis 42 Tage, also rund 3.5 bis 3.8 Prozent aller "
              "Tage.")
    para(doc, "Die Prüfung der auffälligsten Tage zeigt, dass es sich um echte Marktereignisse handelt "
              "und nicht um Datenfehler. Der stärkste Tagesverlust von GBP/USD, minus 4.25 Prozent am "
              "26. September 2022, ist die dokumentierte Marktreaktion auf die Ankündigung umfangreicher "
              "schuldenfinanzierter Steuersenkungen in Grossbritannien vom 23. September 2022. Der "
              "stärkste Tagesverlust von EUR/CHF, minus 1.98 Prozent am 17. Juni 2022, folgt unmittelbar "
              "auf die überraschende Zinserhöhung der Schweizerischen Nationalbank vom 16. Juni 2022.")
    para(doc, "Unser Umgang: Wir behalten alle Ausreisser. Ein Entfernen oder Kappen wäre hier falsch, "
              "denn genau diese starken Bewegungen sind die Reaktion des Markts auf Nachrichten und damit "
              "der Kern unserer Forschungsfrage. Bei Zeitreihen sind extreme Werte oft die wichtigste "
              "Information. Stattdessen begrenzen wir den Einfluss von Ausreissern dort, wo er stören "
              "würde, durch robuste Verfahren: Beim Tages-Sentiment fassen wir mehrere Artikel mit dem "
              "Median statt dem Durchschnitt zusammen (Kapitel 7.1), und beim Quellenvergleich mitteln "
              "wir erst nach der zeitlichen Ausrichtung (Kapitel 6.2).")

    h2(doc, "5.5 Qualitätsprüfung über die Quellen: ein Datierungsfehler")
    para(doc, "Mehrere Quellen sind nur dann ein Gewinn, wenn man sie gegeneinander prüft. Wir haben "
              "deshalb getestet, ob sich zwei Anbieter desselben Wechselkurses in ihren Tagesveränderungen "
              "gleichen. Die Tagesveränderung ist der Unterschied des Kurses von einem Tag zum nächsten. "
              "Stimmen zwei Quellen hier nicht überein, stimmt etwas mit der zeitlichen Zuordnung nicht.")
    add_table(
        doc, "Tabelle 4: Übereinstimmung von Yahoo und EODHD vor der Korrektur",
        ["Wechselkurs", "Übereinstimmung am selben Tag", "Mittlere Differenz", "Bewertung"],
        [
            ["EUR/CHF", "0.92", "rund 2 Pips", "in Ordnung"],
            ["EUR/USD", "0.00", "rund 43 Pips", "fehlerhaft"],
            ["GBP/USD", "0.00", "rund 50 Pips", "fehlerhaft"],
        ],
        widths=[3.0, 5.0, 3.5, 3.0],
    )
    para(doc, "Zwei echte Datenquellen für denselben EUR/USD- oder GBP/USD-Kurs dürften sich am selben Tag "
              "niemals um 43 oder 50 Pips unterscheiden, und ihre Tagesveränderungen müssten fast "
              "vollständig übereinstimmen, nicht bei null liegen (ein Pip ist die kleinste übliche "
              "Kursbewegung, die vierte Nachkommastelle). Verschiebt man die EODHD-Reihe um genau einen "
              "Kalendertag, springt die Übereinstimmung auf 0.66 beziehungsweise 0.89. Die "
              "EODHD-Tagesreihen waren also für EUR/USD und GBP/USD um einen Tag gegenüber Yahoo "
              "versetzt, für EUR/CHF nicht.")
    para(doc, "Die Ursache liegt in unterschiedlichen Konventionen, wie ein Tagesbalken datiert wird. "
              "Yahoo Finance stempelt seine Tageswerte uneinheitlich mal auf 23:00 Uhr des Vortags, mal "
              "auf 00:00 Uhr (jeweils Weltzeit). EODHD ordnet den Sonntags-Balken je nach Paar anders zu. "
              "In Summe ergibt sich der Versatz von einem Tag.")
    para(doc, "Warum das wichtig war: In der Analyse mitteln wir die Quellen pro Tag (Kapitel 7.3). "
              "Mittelt man zwei um einen Tag versetzte Reihen, mischt man zwei verschiedene Markttage zu "
              "einem Wert, und die daraus berechneten Tagesveränderungen werden teils unsinnig.")
    para(doc, "Entscheidend ist die Grössenordnung. Dass sich zwei Forex-Anbieter leicht unterscheiden, "
              "ist normal, denn der Devisenmarkt ist dezentral, es gibt keinen einzigen offiziellen Kurs, "
              "und jeder Handelsplatz quotiert minimal anders. Genau dafür ist ein Mittelwert das richtige "
              "Werkzeug. Bei EUR/CHF betrug der Unterschied rund 2 Pips, also normale Anbieter-Streuung. "
              "Bei EUR/USD waren es rund 43 Pips, das Zwanzigfache. So weit driften echte Kurse nicht "
              "auseinander, das war ein Datierungsfehler. Nach der Korrektur schrumpft der Unterschied auf "
              "rund 12 beziehungsweise 1 Pip. Abbildung 3 zeigt die Übereinstimmung vor und nach der "
              "Korrektur.")
    add_figure(doc, FIGDIR / "fig_source_alignment.png", "Abbildung 3: Quellenvergleich vor und nach der Datums-Ausrichtung. Vor der Korrektur stimmen Yahoo und EODHD bei EUR/USD und GBP/USD kaum überein, danach sehr gut.", width=14)
    para(doc, "Als Nebenbefund war auch MetaTrader bei EUR/USD um einen Tag versetzt. Nach der Korrektur "
              "stimmt es zu 0.96 mit Yahoo überein, was die Quellen gegenseitig bestätigt. Die Behebung "
              "beschreibt Kapitel 6.2.")

    # =====================================================================
    h1(doc, "6. Harmonisierung")
    para(doc, "Ziel der Harmonisierung ist, alle Quellen auf dasselbe Format und denselben Markttag zu "
              "bringen, damit ein Vergleich überhaupt möglich ist.")
    h2(doc, "6.1 Vereinheitlichung von Format und Benennung")
    bullet(doc, "Spaltennamen: Yahoo schreibt Open, High, Low, Close gross, EODHD klein, MetaTrader mit "
                "spitzen Klammern. Wir vereinheitlichen alles auf die Kleinschreibung open, high, low, close.")
    bullet(doc, "Datumsformat: Die Quellen liefern unterschiedliche Schreibweisen des Datums. Wir wandeln "
                "alle in ein einheitliches Datumsformat um.")
    bullet(doc, "Schreibweise der Währungspaare: Yahoo nutzt EURUSD=X, EODHD nutzt EURUSD.FOREX, "
                "MetaTrader nutzt EURUSD. Intern verwenden wir einheitlich EUR_USD, EUR_CHF und GBP_USD.")
    bullet(doc, "Nachrichtenfelder: Den verschachtelten Stimmungswert von EODHD flachen wir in eigene "
                "Spalten auf, damit er wie eine normale Tabellenspalte nutzbar ist.")
    h2(doc, "6.2 Zeitliche Ausrichtung der Quellen")
    para(doc, "Der wichtigste Harmonisierungsschritt ist die Korrektur des in Kapitel 5.5 gefundenen "
              "Datierungsfehlers. Das Skript regenerate_forex_combined.py geht dabei wie folgt vor:")
    bullet(doc, "Yahoo dient als zeitliche Referenz. Seine uneinheitlichen Zeitstempel werden korrekt auf "
                "den jeweils richtigen Handelstag gerundet.")
    bullet(doc, "Für jede andere Quelle wird der Versatz zwischen minus zwei und plus zwei Tagen gesucht, "
                "der die Übereinstimmung der Tagesveränderungen mit Yahoo am grössten macht.")
    bullet(doc, "Ein Versatz wird nur dann angewendet, wenn er die Übereinstimmung deutlich verbessert. So "
                "bleibt EUR/CHF unverändert, während EUR/USD und GBP/USD um einen Tag korrigiert werden.")
    para(doc, "Erst nach dieser Ausrichtung werden die Quellen zusammengeführt und gemittelt.")
    para(doc, "Eine zusätzliche Qualitätsprüfung bestätigt die MetaTrader-Daten von innen heraus: Wir "
              "haben die 15-Minuten-Daten zu Tageswerten zusammengefasst und mit dem Tages-Export "
              "verglichen. An allen 1037 gemeinsamen Tagen stimmen die Schlusskurse exakt überein, ohne "
              "eine einzige Abweichung.")

    # =====================================================================
    h1(doc, "7. Aufbereitung der Analysegrössen")

    h2(doc, "7.1 Tages-Sentiment aus mehreren Artikeln")
    para(doc, "An einem Tag gibt es oft mehrere Artikel mit je einem Stimmungswert. Wir fassen sie pro Tag "
              "zum Median zusammen, also dem mittleren Wert, nicht zum Durchschnitt. Der Hintergrund ist "
              "unsere Annahme zur Tagesstimmung: Wenn etwas Bedeutendes passiert, erscheinen dazu mehrere "
              "Artikel mit deutlichem Sentiment. Ein einzelner oder zwei stark formulierte Artikel deuten "
              "dagegen nicht auf ein grosses Ereignis hin, denn sonst gäbe es mehr Artikel. Der Median "
              "bildet deshalb die breite Stimmung des Tages ab und wird von solchen einzelnen Ausreissern "
              "nicht verzerrt, während der Durchschnitt von einem einzelnen extremen Artikel stark "
              "beeinflusst würde.")
    para(doc, "Ergänzend lässt sich auch der Durchschnitt je Tag bilden. Das ist sinnvoll, wenn man "
              "umgekehrt prüfen möchte, ob einzelne, stark formulierte Artikel eine eigene Bedeutung haben "
              "und nicht untergehen sollen.")

    h2(doc, "7.2 Warum wir Veränderungen und nicht nur das Kursniveau betrachten")
    para(doc, "Für den statistischen Vergleich betrachten wir neben dem Kursniveau vor allem die "
              "Kursveränderung, also den Unterschied von einem Zeitpunkt zum nächsten. Der Grund: Kurse "
              "haben über Jahre einen langfristigen Trend. Vergleicht man zwei Kursniveaus direkt, misst "
              "man vor allem, ob beide demselben Trend folgen, und nicht, ob der Kurs gerade auf eine "
              "Nachricht reagiert. Mit der Veränderung messen wir die eigentliche Reaktion.")
    para(doc, "Die eingangs gestellte anschauliche Frage betrifft aber genau das Niveau: Sinkt der Kurs "
              "nach negativen Nachrichten nachträglich weiter? Diese Niveau-Sicht behandeln wir deshalb in "
              "Kapitel 9.3 ausdrücklich, sowohl mit einer Grafik des Kursverlaufs als auch, indem wir die "
              "kumulative Kursbewegung über die folgenden Wochen messen.")

    h2(doc, "7.3 Mittelwert über die Quellen")
    para(doc, "Wo mehrere Quellen einen Kurs liefern, bilden wir den Mittelwert über die an diesem Tag "
              "vorhandenen Quellen. Wie in Kapitel 5.5 erläutert, hat der Devisenmarkt keinen einzigen "
              "wahren Kurs, und der Mittelwert glättet die normale, kleine Streuung zwischen den Anbietern. "
              "Wichtig ist, dass dieser Mittelwert erst gebildet wird, nachdem die Quellen auf denselben "
              "Markttag ausgerichtet sind.")

    h2(doc, "7.4 Interpolation der Kurse")
    para(doc, "Eine Interpolation füllt fehlende Tage rechnerisch auf. Liegt der Kurs heute bei 1 und "
              "übermorgen bei 3, setzt die lineare Interpolation den fehlenden Tag dazwischen auf 2. Das "
              "macht die Zeitreihe lückenlos und erleichtert spätere Berechnungen. Wir haben beide "
              "Varianten gerechnet und verglichen.")
    add_table(
        doc, "Tabelle 5: Zusammenhang am selben Tag mit und ohne Interpolation",
        ["Wechselkurs", "ohne Interpolation", "mit Interpolation"],
        [["EUR/USD", "0.18", "0.11"], ["GBP/USD", "0.21", "0.16"]],
        widths=[4.0, 5.0, 5.0],
    )
    para(doc, "Mit Interpolation fällt der Zusammenhang. Das ist nachvollziehbar: Die zusätzlich "
              "eingefügten Werte an Wochenenden sind rechnerisch erfunden und enthalten keine neue "
              "Information aus Nachrichten. Sie vergrössern die Datenmenge, ohne ein Signal hinzuzufügen, "
              "und verwässern dadurch den Zusammenhang. Zugleich ist die Interpolation nur eine Näherung, "
              "denn auch Kurse können sich von Tag zu Tag stark ändern, je nachdem was passiert. Deshalb "
              "verwenden wir die Variante ohne Interpolation als Hauptauswertung und die interpolierte "
              "Variante als Vergleich. Beide führen zur selben Kernaussage.")
    para(doc, "Damit ist auch die Frage beantwortet, wie sich der Zusammenhang ohne Interpolation "
              "berechnen lässt: Für den statistischen Vergleich verwenden wir nur die Tage, an denen "
              "sowohl ein Kurs als auch ein Stimmungswert vorliegt. Ein Wert für jeden Kalendertag ist "
              "dafür nicht nötig.")

    h2(doc, "7.5 Zusammenfassung auf andere Zeithorizonte")
    para(doc, "Neben der Tagesebene fassen wir die Daten auf Wochen, Monate und Quartale zusammen. So "
              "lässt sich derselbe Zusammenhang über verschiedene Zeithorizonte prüfen. Fehlende Tage "
              "werden dabei nicht vorher gefüllt, sondern beim Zusammenfassen übersprungen.")

    # =====================================================================
    h1(doc, "8. Sentiment-Analyse")
    para(doc, "Wir bestimmen die Stimmung auf mehreren Wegen, damit sichtbar wird, wie stark das Ergebnis "
              "von der Methode abhängt.")
    h2(doc, "8.1 Hauptweg: das fertige Sentiment von EODHD")
    para(doc, "EODHD liefert pro Artikel bereits einen Stimmungswert zwischen minus eins und plus eins. "
              "Wie dieser Wert genau berechnet wird, legt der Anbieter nicht offen. Wir behandeln ihn "
              "deshalb als gegebenen Wert und nutzen ihn als Hauptgrundlage der Analyse.")
    h2(doc, "8.2 Eigene Berechnung auf dem Nachrichtentext von EODHD")
    para(doc, "Als erste eigene Umsetzung nehmen wir denselben Nachrichtentext von EODHD und berechnen das "
              "Sentiment selbst mit dem Werkzeug TextBlob, statt den fertigen Wert zu übernehmen. So zeigen "
              "wir, dass wir die Stimmungsanalyse selbst durchführen können, und können zugleich "
              "vergleichen, wie nah unsere Berechnung an den Wert von EODHD herankommt. Abbildung 4 stellt "
              "beide Werte für denselben Artikel gegenüber.")
    add_figure(doc, FIGDIR / "fig_sentiment_compare.png", "Abbildung 4: Eigenes Sentiment (TextBlob) im Vergleich zum vorberechneten Sentiment von EODHD auf demselben Artikeltext.", width=12.5)
    para(doc, "Die Übereinstimmung ist mit einer Korrelation von rund 0.27 nur schwach. Ein einfaches, "
              "auf Allgemeinsprache ausgelegtes Verfahren wie TextBlob reproduziert die Werte des "
              "spezialisierten EODHD-Verfahrens also nur teilweise. Das zeigt, wie stark das Ergebnis "
              "einer Sentiment-Analyse von der gewählten Methode abhängt.")
    h2(doc, "8.3 Eigene Nachrichten und eigenes Sentiment über Webscraping")
    para(doc, "Als zweite eigene Umsetzung beschaffen wir die Nachrichten vollständig selbst über "
              "Webscraping (Kapitel 4.4) und berechnen das Sentiment ebenfalls selbst mit TextBlob, aus "
              "Titel und Kurzbeschreibung des Artikels. Auch hier bilden wir pro Tag den Median. Damit "
              "demonstrieren wir den vollen Weg von der eigenen Datenbeschaffung bis zum fertigen "
              "Stimmungswert.")
    h2(doc, "8.4 Grenzen und Abdeckung")
    para(doc, "TextBlob schlägt Wörter in einem allgemeinen Wörterbuch nach und ist nicht auf Finanztexte "
              "spezialisiert. Für fachsprachliche Begriffe liefert es deshalb oft den Wert null, in "
              "unseren Daten bei rund einem Drittel der Artikel. Wir setzen TextBlob dennoch ein, weil es "
              "transparent und reproduzierbar ist. Spezialisierte Verfahren für Finanztexte wären eine "
              "sinnvolle Erweiterung und sind in Kapitel 10 als nächster Schritt genannt.")
    para(doc, "Die Abdeckung der selbst gesammelten Nachrichten ist begrenzt. Insgesamt liegen 127 Tage "
              "mit Artikeln vor. Davon stammt nur ein einziger Tag aus dem September 2024, die eigentliche "
              "Abdeckung beginnt im September 2025 und reicht bis April 2026. An 58 der 127 Tage gibt es "
              "nur einen einzigen Artikel. Diese dünne Grundlage ist der Grund, warum diese zweite "
              "Umsetzung eine Machbarkeitsdemonstration bleibt und keine belastbare Aussage liefert.")

    # =====================================================================
    h1(doc, "9. Analyse und Ergebnisse")

    h2(doc, "9.1 Vorgehen")
    para(doc, "Wir vergleichen das Sentiment eines Zeitraums mit der Kursveränderung desselben und der "
              "folgenden Zeiträume. Auf Tagesebene betrachten wir Verschiebungen von minus zehn bis plus "
              "zehn Tagen. Als Mass dient die Korrelation, eine Zahl zwischen minus eins und plus eins, "
              "die angibt, wie stark zwei Grössen zusammenhängen. Ein Wert nahe null bedeutet keinen "
              "Zusammenhang. Damit ein Wert als gesichert gilt, muss er ein Konfidenzband überschreiten, "
              "das von der Anzahl der Beobachtungen abhängt; bei rund 1100 Beobachtungen liegt es bei etwa "
              "plus minus 0.06. Abbildung 5 zeigt den Zusammenhang für jede Verschiebung.")
    para(doc, "So entsteht die Kurve in Abbildung 5: Für jede Verschiebung k bilden wir alle Paare aus "
              "dem Sentiment eines Tages und der Kursveränderung k Tage später und berechnen über alle "
              "rund 1100 Paare eine einzige Korrelation. Ein Punkt der Kurve ist also kein Mittelwert "
              "über Tage, sondern eine Korrelation über den gesamten Zeitraum. Beispiel Verschiebung "
              "plus 2: Das Sentiment vom Montag wird mit der Kursveränderung vom Mittwoch verglichen, "
              "das Sentiment vom Dienstag mit der vom Donnerstag, und so weiter über alle Tage. Negative "
              "Verschiebungen drehen die Reihenfolge um: Bei minus 2 wird das Sentiment vom Mittwoch mit "
              "der Kursveränderung vom Montag davor verglichen. Das ist kein rückwirkender Einfluss, "
              "sondern der Test der Alternative A2 aus Kapitel 1.3: Wenn die Nachrichten erst über "
              "bereits geschehene Kursbewegungen berichten, müsste der Zusammenhang auf der negativen "
              "Seite hoch sein.")
    add_figure(doc, FIGDIR / "fig_leadlag.png", "Abbildung 5: Zusammenhang zwischen Sentiment und Kursveränderung je zeitlicher Verschiebung auf Tagesebene. Das Maximum liegt bei der Verschiebung null. Die hellen Bänder zeigen den Bereich, in dem ein Wert als zufällig gilt.", width=15.5)

    h2(doc, "9.2 Ergebnis auf Tagesebene")
    add_table(
        doc, "Tabelle 6: Stärkster Zusammenhang je Wechselkurs auf Tagesebene",
        ["Wechselkurs", "stärkste Verschiebung", "Zusammenhang dort", "gleichzeitig", "Beobachtungen"],
        [
            ["EUR/USD", "0 Tage (gleichzeitig)", "0.18", "0.18", "1108"],
            ["GBP/USD", "0 Tage (gleichzeitig)", "0.21", "0.21", "1105"],
            ["EUR/CHF", "nicht aussagekräftig", "nur 9 Tage", "nur 9 Tage", "9"],
        ],
        row_colors=[COL_REAL, COL_REAL, COL_REAL],
        widths=[2.6, 4.2, 3.0, 2.8, 2.8],
    )
    para(doc, "Bei EUR/USD und GBP/USD liegt der stärkste Zusammenhang eindeutig bei der Verschiebung "
              "null, also bei gleichzeitiger Bewegung, mit Werten von 0.18 und 0.21. Diese Werte liegen "
              "klar über dem Konfidenzband und sind damit gesichert. Mit wachsender Verschiebung in beide "
              "Richtungen fällt der Zusammenhang rasch ab. Auffällig ist, dass die Kurve bei der "
              "Verschiebung von minus einem Tag noch leicht über dem Band liegt (EUR/USD 0.10, GBP/USD "
              "0.07): Die Nachrichten eines Tages hängen also auch mit der Kursbewegung des "
              "unmittelbar vorangehenden Tages zusammen. Das passt zur Alternative A2, denn ein Teil der "
              "Berichterstattung folgt dem Markt. Für EUR/CHF liegen nur neun gemeinsame Tage vor, was "
              "für eine Aussage nicht ausreicht.")

    h2(doc, "9.3 Niveau- und Wochen-Sicht: läuft das Sentiment voraus?")
    para(doc, "Die anschauliche Frage lautet: Wenn das Sentiment in einer Woche negativ war, fällt der "
              "Kurs in den darauffolgenden Wochen? Abbildung 6 stellt für EUR/USD je Woche das Kursniveau "
              "(dunkle Linie) und das Nachrichten-Sentiment (Balken, blau positiv, orange negativ) dar. "
              "So lässt sich direkt ablesen, ob auf eine Woche mit negativem Sentiment ein fallender Kurs "
              "folgt. Ein solches Muster ist im Bild nicht erkennbar: Negative Sentiment-Wochen werden "
              "nicht systematisch von fallenden Kursen abgelöst. Weil über vier Jahre viele Wochen "
              "nebeneinander liegen, zeigt der untere Teil jeder Abbildung zusätzlich die letzten zwölf "
              "Monate vergrössert. Dort ist auch die Wochen-Aggregation direkt sichtbar: Jeder Punkt der "
              "Kurslinie ist ein Wochenschlusskurs, die gestrichelte Verbindung zeigt, dass dazwischen "
              "keine weiteren Werte liegen.")
    para(doc, "Zum Lesen der Balken ist eine Unterscheidung wichtig. Eine Woche ohne sichtbaren Balken "
              "ist keine Woche ohne Nachrichten: Auf Wochenebene gibt es in jeder der 225 Wochen "
              "mindestens einen Artikel. Ein fehlender Balken bedeutet stattdessen, dass der Median der "
              "Woche exakt neutral war, also null. Das kommt häufig vor (bei EUR/USD in 126, bei GBP/USD "
              "in 95 der 225 Wochen), weil viele Artikel ein Sentiment von genau null tragen. Diese "
              "neutralen Wochen sind in den Abbildungen als graue Punkte auf der Nulllinie markiert. "
              "Tage ganz ohne Artikel gibt es nur auf Tagesebene (Kapitel 5.2); beim Zusammenfassen zu "
              "Wochen verschwinden diese Lücken.")
    add_figure(doc, FIGDIR / "fig_weekly_EUR_USD.png", "Abbildung 6: EUR/USD, Kursniveau und wöchentliches Nachrichten-Sentiment. Oben der gesamte Zeitraum, unten die letzten zwölf Monate vergrössert. Punkte auf der Kurslinie sind Wochenschlusskurse, graue Punkte auf der Nulllinie sind neutrale Wochen (Median 0).", width=16.5)
    add_figure(doc, FIGDIR / "fig_weekly_GBP_USD.png", "Abbildung 7: GBP/USD, Kursniveau und wöchentliches Nachrichten-Sentiment. Oben der gesamte Zeitraum, unten die letzten zwölf Monate vergrössert. Punkte auf der Kurslinie sind Wochenschlusskurse, graue Punkte auf der Nulllinie sind neutrale Wochen (Median 0).", width=16.5)
    para(doc, "In der vergrösserten Ansicht fallen durchaus einzelne Episoden auf, die nach einem "
              "Vorlauf aussehen. Im März 2026 etwa fällt der Kurs in zwei Wochen mit stark negativem "
              "Sentiment deutlich, und in positiven Wochen steigt er wieder. Solche Episoden sind echt, und sie passen zum Ergebnis des Granger-Tests in Kapitel 9.4, der dem "
              "Sentiment einen kleinen, messbaren Vorhersagebeitrag zuschreibt. Als Beleg für eine "
              "verlässliche Vorhersage taugen sie aber nicht, denn im selben Bild gibt es ebenso Wochen, "
              "in denen das Muster ausbleibt. Wer nur die passenden Episoden herausgreift, überschätzt "
              "den Effekt. Massgebend ist deshalb die Auswertung über alle 225 Wochen.")
    para(doc, "Um das nicht nur visuell, sondern auch in Zahlen zu prüfen, messen wir den Zusammenhang "
              "zwischen dem Sentiment einer Woche und der kumulativen Kursbewegung bis ein, zwei, drei und "
              "vier Wochen später. Wenn das Sentiment vorausliefe, müsste dieser Zusammenhang für die "
              "Folgewochen deutlich von null verschieden sein.")
    add_table(
        doc, "Tabelle 7: Sentiment einer Woche und Kursbewegung in den Folgewochen",
        ["Wechselkurs", "selbe Woche", "bis 1 Woche später", "bis 2 Wochen später", "bis 3 Wochen später"],
        [
            ["EUR/USD", "0.23", "0.07", "0.01", "minus 0.03"],
            ["GBP/USD", "0.17", "0.07", "0.07", "0.01"],
        ],
        row_colors=[COL_REAL, COL_REAL],
        widths=[2.8, 3.0, 3.6, 3.6, 3.6],
    )
    add_figure(doc, FIGDIR / "fig_forward_weeks.png", "Abbildung 8: Sentiment einer Woche und kumulative Kursbewegung der Folgewochen. Der Zusammenhang besteht nur in derselben Woche und fällt danach gegen null.", width=15)
    para(doc, "Das Muster ist klar: Der Zusammenhang ist in derselben Woche am stärksten und fällt danach "
              "rasch gegen null. Auch über das Kursniveau und über mehrere Folgewochen betrachtet sagt das "
              "Sentiment einer Woche die spätere Kursbewegung nicht voraus. Stimmung und Kurs bewegen sich "
              "gemeinsam in derselben Periode.")
    para(doc, "Dieselbe Prüfung haben wir neben der Tages- und Wochenebene auch auf Monatsebene "
              "gemacht. Tabelle 8 stellt den Zusammenhang im selben Zeitraum für alle drei Zeithorizonte "
              "nebeneinander.")
    add_table(
        doc, "Tabelle 8: Zusammenhang im selben Zeitraum je Zeithorizont",
        ["Wechselkurs", "Tag", "Woche", "Monat", "Beobachtungen (Tag / Woche / Monat)"],
        [
            ["EUR/USD", "0.18", "0.23", "0.09", "1108 / 224 / 51"],
            ["GBP/USD", "0.21", "0.17", "0.27", "1105 / 224 / 51"],
        ],
        row_colors=[COL_REAL, COL_REAL],
        widths=[2.8, 2.2, 2.2, 2.2, 6.0],
    )
    para(doc, "Auf Tages- und Wochenebene ist der gleichzeitige Zusammenhang gesichert. Auf Monatsebene "
              "liegen nur 51 Monate vor, und bei so wenigen Beobachtungen gilt ein Wert erst ab etwa "
              "plus minus 0.27 als gesichert. Die Monatswerte von 0.09 und 0.27 erreichen diese Schwelle "
              "nicht oder nur knapp und sind deshalb nur als Tendenz zu lesen. Ein Vorlauf zeigt sich "
              "auch hier nicht: Das Sentiment eines Monats hängt mit der Kursbewegung des Folgemonats "
              "bei beiden Wechselkursen leicht negativ zusammen (minus 0.09 und minus 0.20), was "
              "ebenfalls nicht gesichert ist. Quartale haben mit 17 Beobachtungen zu wenige Datenpunkte "
              "für eine Aussage.")

    h2(doc, "9.4 Formaler Vorhersagetest (Granger-Test)")
    para(doc, "Die Korrelation misst den Zusammenhang je Verschiebung einzeln. Ein strengerer, formaler "
              "Test ist der Granger-Test. Er prüft, ob die Vergangenheit des Sentiments die Vorhersage "
              "der Kursveränderung verbessert, über das hinaus, was die eigene Vergangenheit der "
              "Kursveränderung bereits leistet. Ein kleiner Wert (unter 0.05) in der folgenden Tabelle "
              "bedeutet, dass das Sentiment einen eigenen Vorhersagebeitrag liefert.")
    add_table(
        doc, "Tabelle 9: Granger-Test, p-Werte für den Vorhersagebeitrag des Sentiments",
        ["Verzögerung", "EUR/USD: Sentiment sagt Kurs voraus", "GBP/USD: Sentiment sagt Kurs voraus"],
        [
            ["1 Tag", "0.001", "0.000"],
            ["2 Tage", "0.000", "0.000"],
            ["3 Tage", "0.001", "0.000"],
            ["4 Tage", "0.001", "0.000"],
            ["5 Tage", "0.002", "0.000"],
        ],
        widths=[3.5, 5.5, 5.5],
    )
    para(doc, "Für beide Wechselkurse ist der Vorhersagebeitrag des Sentiments bei allen getesteten "
              "Verzögerungen statistisch signifikant. Bei rund 1100 Beobachtungen werden allerdings schon "
              "sehr kleine Effekte signifikant. Die Stärke dieses Beitrags ist gering, passend zu den "
              "schwachen Korrelationen bei Verschiebungen grösser null. Das Sentiment trägt also ein "
              "schwaches, aber messbares Vorhersagesignal, ist für sich genommen jedoch kein praktisch "
              "brauchbarer Frühindikator.")
    para(doc, "Aufschlussreich ist die Gegenrichtung. Bei GBP/USD verbessert die Vergangenheit des "
              "Kurses die Vorhersage des Sentiments nicht (alle Werte über 0.29); der Zusammenhang läuft "
              "dort also einseitig vom Sentiment zum Kurs. Bei EUR/USD ist der Effekt in beide Richtungen "
              "schwach signifikant. Da Sentiment und Kurs zudem stark gleichzeitig zusammenhängen, kann "
              "ein Teil dieses Vorhersagebeitrags indirekt aus dem gleichzeitigen Zusammenhang stammen.")
    para(doc, "Zwei Einschränkungen gehören zur ehrlichen Einordnung. Erstens zählen unsere "
              "Verzögerungen Kalendertage. Wegen der Wochenenden steckt in der Verzögerung von einem Tag "
              "manchmal auch ein Abstand von drei Tagen, etwa von Freitag zu Montag. Zweitens nimmt der "
              "verwendete Test an, dass die Kurse über die ganze Zeit gleich stark schwanken. Tatsächlich "
              "wechseln Finanzmärkte zwischen ruhigen und unruhigen Phasen, wodurch die p-Werte etwas zu "
              "klein ausfallen können. Beides ändert das Gesamtbild nicht, mahnt aber, die sehr kleinen "
              "Werte nicht überzubewerten.")
    para(doc, "Alle p-Werte beider Richtungen erzeugt das Skript regenerate_lead_lag_results.py "
              "reproduzierbar und legt sie in der Datei granger_results.csv ab; die Tabelle oben zitiert "
              "daraus die Richtung vom Sentiment zur Kursveränderung.")

    h2(doc, "9.5 Warum eine reine Kursgrafik einen Vorlauf vermuten lässt")
    para(doc, "Wenn man Kursniveau und Sentiment beide als Verlauf über die Zeit zeichnet, kann der "
              "Eindruck entstehen, das Sentiment laufe dem Kurs voraus. Dieser Eindruck entsteht, weil "
              "beide Reihen über den langen Zeitraum demselben übergeordneten Trend folgen und sich deshalb "
              "gemeinsam bewegen. Ein echter Vorlauf würde sich als deutlicher Ausschlag bei einer ganz "
              "bestimmten Verschiebung zeigen, nicht als gleichmässiges Mitlaufen. Genau deshalb betrachten "
              "wir zusätzlich die Veränderungen und die kumulative Bewegung in den Folgewochen "
              "(Kapitel 9.3), die den scheinbaren Vorlauf nicht bestätigen.")

    h2(doc, "9.6 Ergebnis der eigenen Datenbeschaffung")
    para(doc, "Auf den selbst über Webscraping gesammelten Nachrichten zeigt sich kein belastbarer "
              "Zusammenhang. Die stärksten Werte erreichen zwar 0.29 und liegen damit über dem bei nur "
              "85 gemeinsamen Tagen sehr breiten Konfidenzband von rund plus minus 0.21. Sie treten aber "
              "bei wechselnden Verschiebungen ohne erkennbares Muster auf, etwa plus neun Tage bei "
              "EUR/USD und plus fünf Tage bei EUR/CHF, und die gleichzeitige Korrelation liegt nahe "
              "null. Hinzu kommt ein Auswahleffekt: Wer den stärksten Wert aus 21 geprüften Verschiebungen herausgreift, findet "
              "bei so wenigen Beobachtungen fast immer einen Wert nahe der Schwelle, auch wenn in "
              "Wahrheit kein Zusammenhang besteht. Das ist angesichts der dünnen Abdeckung (Kapitel 8.4) "
              "zu erwarten. Die eigene Datenbeschaffung lässt sich also vollständig umsetzen, liefert "
              "aufgrund der Datenmenge aber keine belastbare Aussage.")

    h2(doc, "9.7 Öl als möglicher Einflussfaktor")
    para(doc, "Wir haben zusätzlich die Ölpreise WTI und Brent erhoben, weil der Ölpreis häufig als "
              "Einflussfaktor auf Währungen genannt wird. Abbildung 9 im Anhang E zeigt ihren Verlauf. "
              "Um zu prüfen, "
              "ob ein Bezug zu den untersuchten Wechselkursen besteht, haben wir die Tagesveränderungen "
              "von Öl mit denen der Wechselkurse verglichen. Der Zusammenhang ist mit Werten zwischen "
              "minus 0.09 und 0.00 praktisch null. Für EUR/USD und GBP/USD lässt sich also kein messbarer "
              "Einfluss des Ölpreises feststellen. Auch zwischen dem Nachrichten-Sentiment und dem "
              "Ölpreis besteht praktisch kein Zusammenhang (Korrelation rund minus 0.04 bis minus 0.07). "
              "Das ist nachvollziehbar, da sich die von uns geladenen Nachrichten auf die Währungspaare "
              "und nicht auf den Ölmarkt beziehen. Wir behalten Öl deshalb als allgemeinen Kontext und als "
              "zuschaltbare Vergleichsreihe, beziehen es aber nicht in die Beantwortung der "
              "Forschungsfrage ein.")

    # =====================================================================
    h1(doc, "10. Diskussion und Beantwortung der Frage")
    para(doc, "Die Forschungsfrage war, ob die Stimmung von Finanznachrichten einen Einfluss auf "
              "Wechselkurse hat und mit welcher zeitlichen Verzögerung.")
    para(doc, "Unsere Antwort: Der Zusammenhang ist überwiegend gleichzeitig. Bei EUR/USD und GBP/USD "
              "bewegen sich Sentiment und Kurs in derselben Periode gemeinsam, mit einem Zusammenhang von "
              "0.18 beziehungsweise 0.21. Das Bild ist damit klar von der Alternative A1 geprägt, der "
              "gleichzeitigen Bewegung. Einen praktisch nutzbaren Vorlauf konnten wir nicht feststellen, "
              "weder in der einfachen Korrelation noch über das Kursniveau in den Folgewochen.")
    para(doc, "Der formale Granger-Test zeigt allerdings eine Feinheit. Die Vergangenheit des Sentiments "
              "liefert einen statistisch signifikanten, wenn auch sehr kleinen Beitrag zur Vorhersage der "
              "Kursveränderung, bei GBP/USD sogar einseitig vom Sentiment zum Kurs. Die Hypothese H1 ist "
              "damit höchstens in einer sehr schwachen Form gestützt: Es gibt ein messbares Vorlaufsignal, "
              "das aber zu klein ist, um als Frühindikator zu dienen. In der Praxis bleibt das Sentiment "
              "eine Begleitinformation.")
    para(doc, "Bei dieser Deutung ist wichtig, dass ein Zusammenhang keine Ursache beweist. Selbst der "
              "klare gleichzeitige Zusammenhang zeigt nur, dass Markt und Nachrichten gemeinsam reagieren, "
              "nicht warum. Ein gemeinsamer dritter Faktor, etwa eine Entscheidung einer Notenbank, kann "
              "beide zugleich bewegen.")
    para(doc, "Die Grenzen der Auswertung gehören zur Antwort. Für EUR/CHF reichten die Nachrichten von "
              "EODHD nicht aus. Die selbst gesammelten Nachrichten waren zu dünn für eine belastbare "
              "Aussage. TextBlob ist für Finanztexte nur eingeschränkt geeignet.")
    para(doc, "Als nächste Schritte bieten sich an: ein spezialisiertes Verfahren zur Sentiment-Messung "
              "für Finanztexte einzusetzen, um die vielen neutralen Werte zu verringern, sowie die eigene "
              "Datenbeschaffung über einen längeren Zeitraum laufen zu lassen, damit auch der selbst "
              "erstellte Nachrichtenstrom eine belastbare Auswertung erlaubt.")

    # =====================================================================
    h1(doc, "Anhang A: Glossar")
    add_table(
        doc, "Tabelle 10: Glossar der wichtigsten Begriffe",
        ["Begriff", "Bedeutung"],
        [
            ["Wechselkurs (Forex)", "Preis einer Währung ausgedrückt in einer anderen, gehandelt am weltweiten Devisenmarkt."],
            ["Sentiment", "Zahl zwischen minus eins und plus eins, die ausdrückt, wie negativ oder positiv ein Text formuliert ist."],
            ["Polarity", "Bezeichnung von EODHD und TextBlob für den Stimmungswert."],
            ["Kursveränderung (Rendite)", "Veränderung des Kurses von einem Zeitpunkt zum nächsten."],
            ["Median", "Der mittlere Wert einer Reihe, unempfindlich gegenüber einzelnen Ausreissern."],
            ["Korrelation", "Mass für den Zusammenhang zweier Grössen, von minus eins bis plus eins."],
            ["Konfidenzband", "Schwelle, ab der ein gemessener Zusammenhang als gesichert und nicht als Zufall gilt."],
            ["Interpolation", "Rechnerisches Auffüllen fehlender Werte aus den Nachbarwerten."],
            ["Pip", "Kleinste übliche Kursbewegung, die vierte Nachkommastelle eines Wechselkurses."],
            ["REST-Schnittstelle (API)", "Programmierschnittstelle, die Daten über Web-Adressen ausliefert."],
            ["RSS", "Standardformat, über das Webseiten ihre neuesten Artikel bereitstellen."],
            ["Webscraping", "Automatisiertes Sammeln von Inhalten aus dem Internet."],
            ["TextBlob", "Programmbibliothek, die aus Text einen Stimmungswert berechnet."],
            ["EODHD", "Finanzdaten-Anbieter, liefert Kurse und Nachrichten."],
            ["MetaTrader 5", "Handelsplattform eines Brokers, hier als brokernahe Datenquelle genutzt."],
        ],
        widths=[5.0, 11.0],
    )

    # =====================================================================
    h1(doc, "Anhang B: Theoretischer Hintergrund mit Quellen")
    para(doc, "Die folgenden Begriffe aus den Modulunterlagen sind für die getroffenen Entscheidungen "
              "relevant. Sie sind hier gebündelt, damit der Hauptteil ohne Theorie auskommt.")
    h3(doc, "Arten fehlender Werte")
    para(doc, "Man unterscheidet drei Arten fehlender Werte: vollständig zufällig fehlend (MCAR), wenn das "
              "Fehlen mit nichts zusammenhängt; zufällig fehlend in Abhängigkeit anderer beobachteter Werte "
              "(MAR); und nicht zufällig fehlend (MNAR), wenn das Fehlen vom fehlenden Wert selbst abhängt. "
              "Die fehlenden Sentiment-Tage in Kapitel 5.2 sind nicht zufällig fehlend, weil das "
              "Vorhandensein eines Artikels direkt davon abhängt, ob etwas passiert ist. Quelle: "
              "Vorlesungsunterlagen zur Imputation (Woche 2).")
    h3(doc, "Ausreisser-Erkennung und Umgang")
    para(doc, "Zwei gängige Verfahren zur Erkennung: Die z-Wert-Regel markiert Werte, die mehr als drei "
              "Standardabweichungen vom Mittelwert entfernt liegen; sie setzt eine annähernde "
              "Normalverteilung voraus. Die Interquartilsabstand-Regel markiert Werte, die mehr als das "
              "Anderthalbfache des Interquartilsabstands unter dem ersten oder über dem dritten Quartil "
              "liegen; sie kommt ohne Verteilungsannahme aus. Für den Umgang stehen Entfernen, Kappen, "
              "Winsorisieren, Transformieren oder Markieren zur Wahl. Bei Zeitreihen gilt die besondere "
              "Regel, dass extreme Ereignisse oft die wichtigste Information tragen und deshalb gesondert "
              "behandelt statt entfernt werden sollten. Quelle: Vorlesungsunterlagen zur Datenbereinigung "
              "(Woche 2).")
    h3(doc, "Median und robuste Masse")
    para(doc, "Robuste Masse wie Median und Interquartilsabstand werden gegenüber Durchschnitt und "
              "Standardabweichung bevorzugt, wenn Ausreisser vorliegen, weil sie von einzelnen Extremwerten "
              "kaum beeinflusst werden. Quelle: Vorlesungsunterlagen zur Skalierung (Woche 3).")
    h3(doc, "Stationarität")
    para(doc, "Zeitreihen mit langfristigem Trend gelten als nicht stationär. Korrelationen auf solchen "
              "Niveaus messen vor allem die Trendübereinstimmung. Die Betrachtung von Veränderungen statt "
              "Niveaus (Kapitel 7.2) macht die Reihen annähernd stationär.")
    h3(doc, "Harmonisierung")
    para(doc, "Heterogenität zwischen Datenquellen tritt auf drei Ebenen auf: in der Syntax (dem "
              "Dateiformat), in der Struktur (dem Aufbau der Tabellen und Spalten) und in der Semantik "
              "(der Bedeutung der Begriffe und Werte). Unser Projekt betraf alle drei Ebenen: "
              "verschiedene Formate (Kapitel 4), verschiedene Spaltennamen und Symbolschreibweisen "
              "(Kapitel 6.1) und verschiedene Datierungskonventionen desselben Handelstags (Kapitel 6.2). "
              "Das nachträgliche Angleichen bereits erhobener Daten heisst retrospektive Harmonisierung, "
              "ist die in der Praxis häufigste Form und begrenzt zugleich die erreichbare Datenqualität. "
              "Der empfohlene Prozess endet mit einem eigenen Validierungsschritt, in unserem Fall der "
              "Prüfung der Übereinstimmung nach der Ausrichtung (Abbildung 3). Quelle: Cheng et al. "
              "(2024), A General Primer for Data Harmonization, Scientific Data 11, Artikel 152, sowie "
              "Vorlesungsunterlagen zur Datenharmonisierung (Woche 4).")
    h3(doc, "Gestaltung der Abbildungen")
    para(doc, "Die Wahl der Diagrammformen folgt der Regel, dass Menschen Positionen und Längen am "
              "genauesten vergleichen können, Flächen und Farbtöne dagegen deutlich schlechter. Deshalb "
              "verwenden wir Liniendiagramme für Verläufe über die Zeit, Balkendiagramme für den "
              "Vergleich einzelner Werte und Punktwolken für den Vergleich zweier Messverfahren, und "
              "verzichten auf Kreisdiagramme und 3D-Darstellungen. Alle Balkendiagramme beginnen bei "
              "null, damit Längenverhältnisse nicht verzerrt werden. Liniendiagramme von Kursverläufen "
              "nutzen dagegen einen an die Daten angepassten Achsenausschnitt, weil sonst die "
              "Bewegungen, um die es gerade geht, nicht mehr erkennbar wären; die Achsen sind dafür "
              "stets beschriftet. Alle Abbildungen verwenden die farbenblind-sichere Okabe-Ito-Palette "
              "(Blau, Orange, Vermillion) und verzichten bewusst auf die Unterscheidung Rot gegen Grün, "
              "die für rund acht Prozent der Männer schwer lesbar ist. Quelle: Vorlesungsunterlagen zur "
              "Datenvisualisierung (Woche 8).")

    # =====================================================================
    h1(doc, "Anhang C: Reproduzierbarkeit")
    para(doc, "Alle Schritte lassen sich der Reihe nach mit einzelnen Befehlen neu ausführen. Beim Laden "
              "der EODHD-Daten ist das Tageslimit des kostenlosen Tarifs zu beachten.")
    code_lines = [
        "source .venv/bin/activate",
        "",
        "# 1. Rohdaten laden",
        "python src/data_loading/yahoo_loader.py",
        "python src/data_loading/eodhd_loader.py",
        "python src/data_loading/eodhd_news_loader.py",
        "python src/data_loading/webscraping_loader.py",
        "python src/data_loading/oil_loader.py",
        "",
        "# 2. Bereinigen, ausrichten, zusammenführen",
        "python scripts/regenerate_forex_combined.py",
        "python scripts/regenerate_webscraping_sentiment.py",
        "",
        "# 3. Analyse-Ergebnisse und Abbildungen erzeugen",
        "python scripts/regenerate_lead_lag_results.py",
        "python scripts/build_report_figures.py",
        "",
        "# 4. Bericht bauen",
        "python scripts/build_report_docx.py",
    ]
    cp = doc.add_paragraph(); cp.paragraph_format.left_indent = Cm(0.5)
    run = cp.add_run("\n".join(code_lines)); run.font.name = "Consolas"; run.font.size = Pt(9)

    # =====================================================================
    h1(doc, "Anhang D: Interaktives Dashboard")
    para(doc, "Begleitend haben wir eine interaktive Anwendung erstellt (Datei dashboard.py, umgesetzt mit "
              "streamlit). Sie diente vor allem dazu, verschiedene Einstellungen auszuprobieren, ohne jedes "
              "Mal den Code anzupassen oder ein Notebook neu auszuführen. Die für den Bericht relevanten "
              "Auswertungen entstehen in den Notebooks und sind als Abbildungen in den Hauptteil "
              "eingebettet. Das Dashboard erlaubt unter anderem den direkten Vergleich der Quellen, die "
              "Anzeige der fehlenden Tage, die Gegenüberstellung der Preisabweichungen, den Vergleich des "
              "eigenen Sentiments mit dem von EODHD, die interaktive Lead/Lag-Auswertung samt "
              "Granger-Ergebnissen sowie die frei kombinierbare Darstellung von Kurs, Öl und Sentiment.")

    # =====================================================================
    h1(doc, "Anhang E: Ergänzende Abbildungen und Tabellen")
    para(doc, "Verlauf der in Kapitel 9.7 als möglicher Einflussfaktor geprüften Ölpreise. Da sich kein "
              "messbarer Zusammenhang mit den untersuchten Wechselkursen zeigte, steht die Abbildung im "
              "Anhang.")
    add_figure(doc, FIGDIR / "fig_oil.png", "Abbildung 9: Ölpreise WTI und Brent im Zeitverlauf (Tagesschlusskurse; bei rund 1100 Handelstagen wird auf einzelne Punktmarkierungen verzichtet).", width=15)
    para(doc, "Abbildung 10 zeigt die monatliche Datenabdeckung aller Kursquellen als Übersicht: Yahoo "
              "und EODHD decken den gesamten Zeitraum durchgehend ab, wobei EODHD durch die "
              "Sonntagswerte sichtbar mehr Tage je Monat liefert. MetaTrader endet mit dem Export im "
              "Dezember 2025.")
    add_figure(doc, FIGDIR / "fig_coverage.png", "Abbildung 10: Monatliche Datenabdeckung je Quelle und Währungspaar (Anzahl Tage mit Daten je Monat).", width=16.5)
    para(doc, "Tabelle 11 zeigt, wie dicht die Nachrichtenlage je Währungspaar ist. Sie macht die in "
              "Kapitel 5.2 beschriebene Schwäche von EUR/CHF greifbar: Während EUR/USD im Mittel rund "
              "19 Artikel pro Tag erhält, sind es bei EUR/CHF insgesamt 12 Artikel mit Stimmungswert im "
              "ganzen Zeitraum.")
    add_table(
        doc, "Tabelle 11: Artikel pro Tag und Währungspaar bei EODHD (nur Artikel mit Stimmungswert)",
        ["Währungspaar", "Artikel", "Tage mit Artikeln", "Mittel pro Tag", "Median pro Tag", "Maximum an einem Tag"],
        [
            ["EUR/USD", "27937", "1440", "19.4", "21", "61"],
            ["GBP/USD", "18044", "1409", "12.8", "14", "37"],
            ["EUR/CHF", "12", "10", "1.2", "1", "3"],
        ],
        row_colors=[COL_REAL, COL_REAL, COL_REAL],
        widths=[3.0, 2.4, 3.2, 2.8, 2.8, 3.4],
    )

    update_fields_on_open(doc)
    doc.save(DST)
    print(f"Gespeichert: {DST} ({DST.stat().st_size/1024:.1f} KB)")


if __name__ == "__main__":
    build()
